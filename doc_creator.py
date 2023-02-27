from word_processor import WordProcessor, DocxEditor
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt

WordProcessor = WordProcessor()
DocxEditor = DocxEditor()


class DocCreator:
    def __init__(self, isUrgent, indexList, placeholderList, file_name, new_content):

        self.dest_doc = WordProcessor.getDocx('')
        self.index = WordProcessor.getDocx('sampleTemplate/index.docx')
        self.urgent = WordProcessor.getDocx('sampleTemplate/urgent.docx')
        self.notice_of_motion = WordProcessor.getDocx(
            'sampleTemplate/notice_of_motion.docx')

        self.memo_of_parties = WordProcessor.getDocx(
            'sampleTemplate/memo_of_parties.docx')

        self.pet_synopsis = WordProcessor.getDocx(
            'sampleTemplate/pet_synopsis.docx')

        self.pet_title = WordProcessor.getDocx(
            'sampleTemplate/pet_title.docx')

        self.new_content = WordProcessor.getDocx(
            'sampleTemplate/newContent.docx')

        # self.vakalatnama = WordProcessor.getDocx(
        #     'sampleTemplate/vakalatnama.docx')

        table_index = self.index.tables[1]
        for i, item in enumerate(indexList, start=1):
            DocxEditor.add_index_table_row(table_index, [str(i), item, " "])
        # for item in indexList:
        #     DocxEditor.add_index_table_row(table_index, ["\u2022", item, " "])

        WordProcessor.addToDocx(self.index, self.dest_doc)

        if isUrgent:
            WordProcessor.addPageBreak(self.dest_doc)  # page break
            WordProcessor.addToDocx(self.urgent, self.dest_doc)

        WordProcessor.addPageBreak(self.dest_doc)  # page break

        WordProcessor.addToDocx(self.notice_of_motion, self.dest_doc)

        WordProcessor.addPageBreak(self.dest_doc)  # page break

        WordProcessor.addToDocx(self.memo_of_parties, self.dest_doc)

        WordProcessor.addPageBreak(self.dest_doc)  # page break

        WordProcessor.addToDocx(self.pet_synopsis, self.dest_doc)

        WordProcessor.addPageBreak(self.dest_doc)  # page break

        WordProcessor.addToDocx(self.pet_title, self.dest_doc)

        if new_content:
            for item in new_content:
                WordProcessor.addPageBreak(self.dest_doc)  # page break
                new_doc = WordProcessor.getDocx(
                    'sampleTemplate/newContent.docx')

                soup = BeautifulSoup(item["richText"], 'html.parser')

                for p in soup.find_all('p'):
                    paragraph = new_doc.add_paragraph()
                    paragraph.add_run(p.get_text())
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.runs[0].font.name = 'Bookman Old Style'
                    paragraph.runs[0].font.size = Pt(13)

                for table in soup.find_all('table'):

                    # Add table to document with desired properties
                    new_table = new_doc.add_table(rows=1, cols=2)
                    new_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Iterate over rows in the HTML table
                    for row in table.find_all('tr'):
                        table_row = new_table.add_row()

                        # Iterate over cells in the HTML row
                        for cell in row.find_all('td'):
                            cell_text = cell.get_text()

                            # Add cell to row with desired properties
                            cell_paragraph = table_row.cells[-1].add_paragraph(
                                cell_text)
                            cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cell_paragraph.runs[0].font.name = 'Bookman Old Style'
                            cell_paragraph.runs[0].font.size = Pt(13)

                for paragraph in new_doc.paragraphs:
                    if '{NAME}' in paragraph.text:
                        # Replace placeholder with desired string
                        paragraph.text = paragraph.text.replace(
                            '{NAME}', item['name'])

                WordProcessor.addToDocx(new_doc, self.dest_doc)

        # WordProcessor.addPageBreak(self.dest_doc)  # page break

        # WordProcessor.addToDocx(self.vakalatnama, self.dest_doc)

        DocxEditor.replace_placeholders(self.dest_doc, placeholderList)

        WordProcessor.saveTolocal(self.dest_doc, file_name)


class Single_docx:
    def __init__(self):
        self.dest_doc = WordProcessor.getDocx('')
        self.index = WordProcessor.getDocx('sampleTemplate/index.docx')
        self.urgent = WordProcessor.getDocx('sampleTemplate/urgent.docx')
        self.notice_of_motion = WordProcessor.getDocx(
            'sampleTemplate/notice_of_motion.docx')

        self.memo_of_parties = WordProcessor.getDocx(
            'sampleTemplate/memo_of_parties.docx')

        self.pet_synopsis = WordProcessor.getDocx(
            'sampleTemplate/pet_synopsis.docx')

        self.pet_title = WordProcessor.getDocx(
            'sampleTemplate/pet_title.docx')

    def get_docx(self, docName, placeholderList):
        if docName == 'Urgent Application':
            WordProcessor.addToDocx(self.urgent, self.dest_doc)
            DocxEditor.replace_placeholders(self.dest_doc, placeholderList)
            return WordProcessor.getBase64(self.dest_doc)

        if docName == 'Notice of Motion':
            WordProcessor.addToDocx(self.notice_of_motion, self.dest_doc)
            DocxEditor.replace_placeholders(self.dest_doc, placeholderList)
            return WordProcessor.getBase64(self.dest_doc)

        if docName == 'Memo of Parties':
            WordProcessor.addToDocx(self.memo_of_parties, self.dest_doc)
            DocxEditor.replace_placeholders(self.dest_doc, placeholderList)
            return WordProcessor.getBase64(self.dest_doc)
